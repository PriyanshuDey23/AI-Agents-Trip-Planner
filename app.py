import streamlit as st
from WorkFlow.Agents import TripCrew, display_trip_results
from WorkFlow.Agentic_Rag import setup_rag_system
from io import BytesIO
from docx import Document as DocxDocument
from datetime import datetime

# Page config and title
st.set_page_config(page_title="AI Travel Planner", layout="wide")
st.title("🛫 AI Travel Planning Assistant")

# Sidebar configuration
st.sidebar.header("🗺️ Trip Preferences")
travel_type = st.sidebar.selectbox("Travel Type", ["Leisure", "Business", "Adventure", "Cultural"])
interests = st.sidebar.multiselect("Interests", ["History", "Food", "Nature", "Art", "Shopping", "Nightlife"])
season = st.sidebar.selectbox("Season", ["Summer", "Winter", "Spring", "Fall"])
duration = st.sidebar.slider("Trip Duration (days)", 1, 14, 7)
budget = st.sidebar.selectbox("Budget Range", ["$500-$1000", "$1000-$2000", "$2000-$5000", "Luxury"])

# Session state
if "trip_result" not in st.session_state:
    st.session_state["trip_result"] = None
if "full_trip_text" not in st.session_state:
    st.session_state["full_trip_text"] = ""
if "timestamp" not in st.session_state:
    st.session_state["timestamp"] = datetime.now().strftime("%Y%m%d_%H%M%S")
if "rag_answer" not in st.session_state:
    st.session_state["rag_answer"] = ""


# Generate Travel Plan button
if st.button("🛫 Generate Travel Plan"):
    with st.spinner("🌎 Planning your trip..."):
        inputs = {
            "travel_type": travel_type,
            "interests": interests,
            "season": season,
            "duration": duration,
            "budget": budget
        }
        try:
            result = TripCrew(inputs).run()
            trip_result_dict = display_trip_results(result)
            st.session_state["trip_result"] = trip_result_dict

            # Create concatenated trip text
            full_trip_text = "\n\n".join(
                f"## {key.replace('_', ' ').title()}\n\n{value}"
                for key, value in trip_result_dict.items()
            )
            st.session_state["full_trip_text"] = full_trip_text
            st.session_state["timestamp"] = datetime.now().strftime("%Y%m%d_%H%M%S")
            st.session_state["chat_history"] = []  # Reset chat when a new trip is planned
            st.success("✅ Trip planned successfully!")
        except Exception as e:
            st.error(f"❌ Trip planning failed: {e}")

# Show trip plan in expander only
if st.session_state["full_trip_text"]:
    with st.expander("📝 View Full Trip Plan", expanded=True):
        st.markdown(
            f"<h3 style='color:#4CAF50;'>🗺️ Your Complete Travel Itinerary 📖</h3>",
            unsafe_allow_html=True
        )

        # Render each section title and content
        for key, content in st.session_state["trip_result"].items():
            section_title = key.replace("_", " ").title()
            st.subheader(f"📌 {section_title}")
            st.markdown(content)

        # Download options
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="📥 Download Itinerary (Markdown)",
                data=st.session_state["full_trip_text"],
                file_name=f"Trip_Itinerary_{st.session_state['timestamp']}.md",
                mime="text/markdown"
            )

        with col2:
            itinerary_doc = DocxDocument()
            itinerary_doc.add_heading("Your Complete Trip Plan", 0)
            for line in st.session_state["full_trip_text"].splitlines():
                itinerary_doc.add_paragraph(line)
            buffer = BytesIO()
            itinerary_doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="📥 Download Itinerary (.docx)",
                data=buffer,
                file_name=f"Trip_Itinerary_{st.session_state['timestamp']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# RAG-based interactive chat if trip is generated
if st.session_state["full_trip_text"]:
    st.markdown("---")
    st.subheader("💬 Ask Questions About Your Itinerary")

    # Initialize chat history in session state
    if "chat_history" not in st.session_state:
        st.session_state["chat_history"] = []

    # Display past chat messages
    for i, (question, answer) in enumerate(st.session_state["chat_history"], 1):
        with st.chat_message("user"):
            st.markdown(f"**{question}**")
        with st.chat_message("assistant"):
            st.markdown(answer)

    # User input via chat_input
    user_query = st.chat_input("Type your question and hit Enter")

    if user_query:
        with st.spinner("🤖 Thinking..."):
            try:
                # Use your agentic RAG setup function on the itinerary content
                response = setup_rag_system(
                    content=st.session_state["full_trip_text"],
                    question=user_query
                )

                # Save the new message to chat history
                st.session_state["chat_history"].append((user_query, response))

                # Immediately display latest interaction
                with st.chat_message("user"):
                    st.markdown(f"**{user_query}**")
                with st.chat_message("assistant"):
                    st.markdown(response)

            except Exception as e:
                st.error(f"❌ Could not get an answer: {e}")

    # Download chat history as Markdown
    if st.session_state["chat_history"]:
        chat_log = "\n\n".join(
            [f"**Q{i+1}:** {q}\n**A{i+1}:** {a}" for i, (q, a) in enumerate(st.session_state["chat_history"])]
        )

        st.download_button(
            label="📥 Download Chat History as Markdown",
            data=chat_log,
            file_name=f"Trip_QA_Chat_{st.session_state['timestamp']}.md",
            mime="text/markdown"
        )
else:
    st.info("📝 Generate your travel plan first to ask questions about it.")


